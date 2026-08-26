from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = "/Users/jennitang/Desktop/Medicine/CV_Updated/CV_JenniferTang_2026.docx"
OUT = "/Users/jennitang/Developer/chalk-talk/CV_JenniferTang_2026_UCSF.docx"

src = Document(SRC)
doc = Document()

NAVY = RGBColor(22, 62, 96)
GRAY = RGBColor(89, 89, 89)
BLACK = RGBColor(0, 0, 0)
FONT = "Arial"

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(.68)
section.bottom_margin = Inches(.68)
section.left_margin = Inches(.72)
section.right_margin = Inches(.72)
section.header_distance = Inches(.28)
section.footer_distance = Inches(.3)
section.different_first_page_header_footer = True

styles = doc.styles

def set_font(style, size, bold=False, color=BLACK):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)

normal = styles["Normal"]
set_font(normal, 9.7)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(2.5)
normal.paragraph_format.line_spacing = 1.04
normal.paragraph_format.widow_control = True
normal.paragraph_format.keep_with_next = False

def make_style(name, size, bold=False, color=BLACK, before=0, after=0, keep=True):
    s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_font(s, size, bold, color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.keep_with_next = keep
    s.paragraph_format.widow_control = True
    return s

name_style = make_style("CV Name", 20, True, NAVY, 0, 1)
tag_style = make_style("CV Tagline", 11, False, GRAY, 0, 2)
contact_style = make_style("CV Contact", 9.2, False, GRAY, 0, 8)
section_style = make_style("CV Section", 11.3, True, NAVY, 8, 3)
subhead_style = make_style("CV Subheading", 10.1, True, BLACK, 5, 2)
entry_style = make_style("CV Entry", 9.7, False, BLACK, 2.5, 0)
detail_style = make_style("CV Detail", 9.7, False, BLACK, 0, 2, keep=False)
detail_style.paragraph_format.left_indent = Inches(.66)
detail_style.paragraph_format.first_line_indent = Inches(0)
list_style = make_style("CV List Entry", 9.55, False, BLACK, 0, 2, keep=False)
list_style.paragraph_format.left_indent = Inches(.20)
list_style.paragraph_format.first_line_indent = Inches(-.20)

def bottom_border(paragraph, color="AAB7C4", size="6", space="2"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)

def copy_runs(source_p, target_p):
    for sr in source_p.runs:
        if not sr.text:
            continue
        r = target_p.add_run(sr.text)
        r.bold = sr.bold
        r.italic = sr.italic
        r.underline = sr.underline
        r.font.name = FONT
        r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        if sr.font.color.rgb:
            r.font.color.rgb = sr.font.color.rgb

def add_text(text, style, bold_prefix_date=False):
    p = doc.add_paragraph(style=style)
    if bold_prefix_date and text and text[0].isdigit():
        parts = text.split(" ", 1)
        r = p.add_run(parts[0] + (" " if len(parts) > 1 else ""))
        r.bold = True
        if len(parts) > 1:
            p.add_run(parts[1])
    else:
        p.add_run(text)
    return p

section_indices = {3, 5, 13, 29, 35, 41, 48, 70, 72, 76, 79, 83, 89, 91, 93}
subhead_indices = {42, 46, 49, 61}
entry_indices = {6, 10, 14, 17, 19, 21, 23, 25, 27, 30, 33, 84, 86, 87}
detail_indices = {7, 8, 9, 11, 12, 15, 16, 18, 20, 22, 24, 26, 28, 31, 32, 34, 85, 88}
list_indices = set(range(36, 41)) | set(range(43, 46)) | set(range(47, 48)) | set(range(50, 72)) | set(range(73, 84)) | set(range(90, 95))

# Entries that were accidentally concatenated in the source are separated here;
# wording is preserved exactly apart from normalizing doubled spaces.
split_entries = {
    39: [
        "2023 Physician and Surgeon, New York State Education Department, G842LNMB — active",
        "2023 Controlled Substance Registration, Drug Enforcement Administration",
    ],
    44: [
        "“Hyponatremia.” NYU Langone Hospital–Brooklyn, Internal Medicine Residency Noon Conference. Brooklyn, NY.",
        "“CKD Management.” NYU Langone Hospital–Brooklyn, Internal Medicine Residency Noon Conference. Brooklyn, NY.",
        "“ESRD and Renal Replacement Therapy.” NYU Langone Hospital–Brooklyn, Internal Medicine Residency Noon Conference. Brooklyn, NY.",
    ],
    51: [
        "“Dialyzer Reactions.” Zuckerberg San Francisco General Hospital Nephrology Conference. San Francisco, CA, 2022.",
        "“Left Ventricular Assist Device and Hemodialysis — A Challenge in Homeostasis.” UCSF Nephrology Case Conference. San Francisco, CA, 2022.",
    ],
    58: [
        "“The Effect of Using Vitamin C, Hydrocortisone, and Thiamine Triple Therapy in the Treatment of Septic Shock.” CHEST Conference on Lung Health, Innovations in Pulmonary Research Symposium. Los Angeles, CA, 2019.",
        "“The Effect of Using Vitamin C, Hydrocortisone, and Thiamine Triple Therapy in the Treatment of Septic Shock.” CHEST Annual Meeting. New Orleans, LA, 2019.",
    ],
    67: [
        "“A Rare Case of Culture-Negative Bartonella Endocarditis Resulting in Septic Coronary Embolism Presenting as ST-Segment Elevation Myocardial Infarction.” American Heart Association Scientific Sessions. Philadelphia, PA, 2019.",
        "“Hypereosinophilia Presenting as Stroke and Acute Coronary Syndrome.” American College of Physicians Internal Medicine Meeting. San Diego, CA, 2019.",
    ],
    74: [
        "2023–2024 Regional Pharmacy & Therapeutics Committee, The Permanente Medical Group",
        "2020–2021 Elected Residency Representative, SAFE Committee, LAC+USC Medical Center",
        "2015–2016 Co-President, Medical Students for Choice, Keck School of Medicine of USC",
    ],
}

for i, sp in enumerate(src.paragraphs):
    if not sp.text.strip():
        continue
    if i == 0:
        p = doc.add_paragraph(style=name_style)
        copy_runs(sp, p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 1:
        p = doc.add_paragraph(style=tag_style)
        copy_runs(sp, p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2:
        p = doc.add_paragraph(style=contact_style)
        copy_runs(sp, p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom_border(p, "AAB7C4", "6", "5")
    elif i in section_indices:
        p = doc.add_paragraph(style=section_style)
        copy_runs(sp, p)
        bottom_border(p)
    elif i in subhead_indices:
        p = doc.add_paragraph(style=subhead_style)
        copy_runs(sp, p)
    elif i in split_entries:
        for text in split_entries[i]:
            add_text(text, list_style, bold_prefix_date=i in {39, 74})
    elif i in entry_indices:
        p = doc.add_paragraph(style=entry_style)
        copy_runs(sp, p)
        p.paragraph_format.keep_with_next = True
    elif i in detail_indices:
        p = doc.add_paragraph(style=detail_style)
        copy_runs(sp, p)
    elif i in list_indices:
        p = doc.add_paragraph(style=list_style)
        copy_runs(sp, p)
    else:
        p = doc.add_paragraph(style=normal)
        copy_runs(sp, p)

# Running header and page number (suppressed on page 1).
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("JENNIFER TANG, MD  |  CURRICULUM VITAE")
hr.font.name = FONT
hr.font.size = Pt(8)
hr.font.bold = True
hr.font.color.rgb = GRAY
bottom_border(hp, "D4DAE0", "4", "2")

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Page ")
fr.font.name = FONT
fr.font.size = Pt(8)
fr.font.color.rgb = GRAY
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Add core metadata and request field refresh in Word.
doc.core_properties.title = "Curriculum Vitae — Jennifer Tang, MD"
doc.core_properties.subject = "Academic Medicine Curriculum Vitae"
doc.core_properties.keywords = "curriculum vitae, nephrology, academic medicine"
settings = doc.settings._element
update = OxmlElement("w:updateFields")
update.set(qn("w:val"), "true")
settings.append(update)

doc.save(OUT)
print(OUT)
